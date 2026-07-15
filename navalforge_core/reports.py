"""Multi-format preliminary report exports."""

from __future__ import annotations

import csv
import json
from pathlib import Path
from typing import Any

from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from .models import EvaluationResult, Project

DISCLAIMER = (
    "Preliminary engineering result. It does not replace the responsible engineer, "
    "formal approval, validated software, detailed structural analysis or verified input data."
)


def _result_dict(result: EvaluationResult | dict[str, Any]) -> dict[str, Any]:
    return result.model_dump(mode="json") if isinstance(result, EvaluationResult) else result


def export_json(result: EvaluationResult | dict[str, Any], path: Path) -> Path:
    path.write_text(json.dumps(_result_dict(result), indent=2, ensure_ascii=False), encoding="utf-8")
    return path


def export_csv(result: EvaluationResult | dict[str, Any], path: Path) -> Path:
    data = _result_dict(result)
    with path.open("w", encoding="utf-8", newline="") as stream:
        writer = csv.writer(stream)
        writer.writerow(["requirement_id", "description", "kind", "required", "actual", "unit", "passed"])
        for row in data["requirements"]["matrix"]:
            writer.writerow(
                [
                    row["id"],
                    row["description"],
                    row["kind"],
                    row["required"],
                    row["actual"],
                    row["unit"],
                    row["passed"],
                ]
            )
    return path


def export_xlsx(result: EvaluationResult | dict[str, Any], path: Path) -> Path:
    data = _result_dict(result)
    book = Workbook()
    summary = book.active
    summary.title = "Summary"
    summary.append(["NavalForge Concept", data["project_id"]])
    summary.append(["Revision", data["revision"]])
    summary.append(["Status", data["status"]])
    summary.append(["Notice", DISCLAIMER])
    summary.append([])
    summary.append(["Indicator", "Value"])
    for key, value in data["indicators"].items():
        if isinstance(value, str | int | float | bool) or value is None:
            summary.append([key, value])

    requirements = book.create_sheet("Requirements")
    headers = ["id", "description", "kind", "metric", "operator", "required", "actual", "unit", "passed"]
    requirements.append(headers)
    for row in data["requirements"]["matrix"]:
        requirements.append([row.get(header) for header in headers])

    weights = book.create_sheet("Weights")
    weight_headers = ["id", "description", "group", "total_weight_kg", "longitudinal_moment_kg_m", "transverse_moment_kg_m", "vertical_moment_kg_m"]
    weights.append(weight_headers)
    for row in data["results"]["weights"]["items"]:
        weights.append([row.get(header) for header in weight_headers])

    hydro = book.create_sheet("Hydrostatics")
    hydro.append(["Parameter", "Value"])
    for key, value in data["results"]["hydrostatics"].items():
        hydro.append([key, value])

    variants = book.create_sheet("Variants")
    variant_headers = ["variant_id", "status", "adherence_percent", "displacement_kg", "speed_kn", "range_nm", "gm_m", "cost_brl", "technical_risk"]
    variants.append(variant_headers)
    for row in data.get("variants", []):
        variants.append([row.get(header) for header in variant_headers])

    for sheet in book.worksheets:
        for cell in sheet[1]:
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill("solid", fgColor="0B5363")
        sheet.freeze_panes = "A2"
        for column in sheet.columns:
            width = min(55, max(12, max(len(str(cell.value or "")) for cell in column) + 2))
            sheet.column_dimensions[column[0].column_letter].width = width
            for cell in column:
                cell.alignment = Alignment(vertical="top", wrap_text=True)
    book.save(path)
    return path


def export_docx(
    project: Project,
    result: EvaluationResult | dict[str, Any],
    path: Path,
) -> Path:
    data = _result_dict(result)
    document = Document()
    document.add_heading("NavalForge Concept", 0)
    document.add_paragraph("Preliminary conceptual design report")
    document.add_heading(project.name, level=1)
    document.add_paragraph(f"Project: {project.project_id} | Revision: {project.revision}")
    notice = document.add_paragraph(DISCLAIMER)
    notice.style = document.styles["Intense Quote"]
    document.add_heading("Engineering gate", level=1)
    document.add_paragraph(data["status"])
    document.add_heading("Main indicators", level=1)
    table = document.add_table(rows=1, cols=3)
    table.style = "Light Shading Accent 1"
    table.rows[0].cells[0].text = "Indicator"
    table.rows[0].cells[1].text = "Value"
    table.rows[0].cells[2].text = "Unit"
    indicators = data["indicators"]
    rows = [
        ("Displacement", indicators["displacement_kg"], "kg"),
        ("Maximum speed", indicators["max_speed_kn"], "kn"),
        ("Required power", indicators["required_power_kw"], "kW"),
        ("Range", indicators["range_nm"], "nmi"),
        ("Corrected GM", indicators["gm_m"], "m"),
        ("Freeboard", indicators["freeboard_m"], "m"),
        ("Adherence", indicators["adherence_percent"], "%"),
    ]
    for label, value, unit in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = f"{float(value):,.2f}"
        cells[2].text = unit

    document.add_heading("Requirement matrix", level=1)
    req_table = document.add_table(rows=1, cols=4)
    req_table.style = "Light Grid Accent 1"
    for cell, text in zip(req_table.rows[0].cells, ["ID", "Requirement", "Actual", "Status"], strict=True):
        cell.text = text
    for row in data["requirements"]["matrix"]:
        cells = req_table.add_row().cells
        cells[0].text = str(row["id"])
        cells[1].text = str(row["description"])
        cells[2].text = f"{row['actual']} {row['unit']}"
        cells[3].text = "PASS" if row["passed"] is True else "FAIL/REVIEW"

    document.add_heading("Active engineering reservations", level=1)
    for warning in data["warnings"]:
        document.add_paragraph(str(warning), style="List Bullet")
    document.add_heading("Selected alternatives", level=1)
    for key, alternative in data.get("selected_alternatives", {}).items():
        if not isinstance(alternative, dict) or not alternative.get("variant_id"):
            continue
        document.add_heading(f"NF-{key.upper()}: {alternative['variant_id']}", level=2)
        document.add_paragraph(str(alternative.get("rationale", "")))
    document.save(path)
    return path


def export_pdf(
    project: Project,
    result: EvaluationResult | dict[str, Any],
    path: Path,
) -> Path:
    data = _result_dict(result)
    styles = getSampleStyleSheet()
    doc = SimpleDocTemplate(
        str(path),
        pagesize=A4,
        rightMargin=16 * mm,
        leftMargin=16 * mm,
        topMargin=16 * mm,
        bottomMargin=16 * mm,
        title=f"NavalForge Concept - {project.name}",
    )
    story: list[Any] = [
        Paragraph("NavalForge Concept", styles["Title"]),
        Paragraph("Preliminary conceptual design report", styles["Heading2"]),
        Spacer(1, 4 * mm),
        Paragraph(f"<b>{project.name}</b> — {project.project_id} — Revision {project.revision}", styles["BodyText"]),
        Spacer(1, 4 * mm),
        Paragraph(f"<b>Engineering gate:</b> {data['status']}", styles["BodyText"]),
        Spacer(1, 3 * mm),
        Paragraph(f"<b>Technical notice:</b> {DISCLAIMER}", styles["BodyText"]),
        Spacer(1, 6 * mm),
        Paragraph("Main indicators", styles["Heading1"]),
    ]
    indicators = data["indicators"]
    indicator_rows = [
        ["Indicator", "Value", "Unit"],
        ["Adherence", f"{indicators['adherence_percent']:.1f}", "%"],
        ["Displacement", f"{indicators['displacement_kg']:.2f}", "kg"],
        ["Maximum speed", f"{indicators['max_speed_kn']:.2f}", "kn"],
        ["Required installed power", f"{indicators['required_power_kw']:.2f}", "kW"],
        ["Range", f"{indicators['range_nm']:.2f}", "nmi"],
        ["Corrected GM", f"{indicators['gm_m']:.3f}", "m"],
        ["Freeboard", f"{indicators['freeboard_m']:.3f}", "m"],
    ]
    indicator_table = Table(indicator_rows, colWidths=[90 * mm, 38 * mm, 25 * mm])
    indicator_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0B5363")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("GRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#8BA5AC")),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#EAF4F5")]),
            ]
        )
    )
    story.extend([indicator_table, Spacer(1, 6 * mm), Paragraph("Requirement matrix", styles["Heading1"])])
    req_rows = [["ID", "Requirement", "Actual", "Result"]]
    for row in data["requirements"]["matrix"]:
        req_rows.append(
            [
                row["id"],
                Paragraph(str(row["description"]), styles["BodyText"]),
                f"{row['actual']} {row['unit']}",
                "PASS" if row["passed"] is True else "FAIL/REVIEW",
            ]
        )
    req_table = Table(req_rows, colWidths=[25 * mm, 75 * mm, 35 * mm, 28 * mm], repeatRows=1)
    req_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0B5363")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("GRID", (0, 0), (-1, -1), 0.3, colors.grey),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
            ]
        )
    )
    story.extend([req_table, PageBreak(), Paragraph("Warnings and assumptions", styles["Heading1"])])
    for warning in data["warnings"]:
        story.append(Paragraph(f"• {warning}", styles["BodyText"]))
    story.append(Spacer(1, 5 * mm))
    for assumption in data["assumptions"]:
        story.append(Paragraph(f"• {assumption}", styles["BodyText"]))
    story.extend([Spacer(1, 7 * mm), Paragraph("Selected alternatives", styles["Heading1"])])
    for key, alternative in data.get("selected_alternatives", {}).items():
        if isinstance(alternative, dict) and alternative.get("variant_id"):
            story.append(Paragraph(f"<b>NF-{key.upper()}</b> — {alternative['variant_id']}", styles["Heading2"]))
            story.append(Paragraph(str(alternative.get("rationale", "")), styles["BodyText"]))
    doc.build(story)
    return path


def generate_report_bundle(
    project: Project,
    result: EvaluationResult,
    output_dir: Path,
) -> list[Path]:
    output_dir.mkdir(parents=True, exist_ok=True)
    stem = project.project_id.lower().replace(" ", "_")
    return [
        export_json(result, output_dir / f"{stem}.json"),
        export_csv(result, output_dir / f"{stem}.csv"),
        export_xlsx(result, output_dir / f"{stem}.xlsx"),
        export_docx(project, result, output_dir / f"{stem}.docx"),
        export_pdf(project, result, output_dir / f"{stem}.pdf"),
    ]
